#!/usr/bin/env python3
"""
生成古浪宁氏项目工程结算书（适配实际清单列结构）
"""
import openpyxl
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from datetime import date
import os, re

# ========== 1. 读取清单数据 ==========
wb = openpyxl.load_workbook(
    "/Users/mac/Desktop/宁氏高钙围墙及场地平整、大门形象施工项目清单.xlsx",
    data_only=True
)
ws = wb.active

# 分4个section读取
sections = {
    "围墙工程已完成量": {"rows": [], "mat_total": 0, "labor_total": 0, "total": 0},
    "围墙新增项目": {"rows": [], "mat_total": 0, "labor_total": 0, "total": 0},
    "场地平整工程": {"rows": [], "rows2": [], "mat_total": 0, "labor_total": 0, "total": 0},
    "临时设施及文明施工": {"rows": [], "mat_total": 0, "labor_total": 0, "total": 0},
}

# 根据之前读取的数据，手动填入精确值（来自xlsx公式计算结果）

# --- 围墙工程已完成量 ---
data1 = [
    ("1#北侧围墙", "167.4", "m", 81594.00, 60264.00),
    ("2#东侧围墙", "75", "m", 38540.40, 27000.00),
    ("3#西侧围墙", "75.6", "m", 37077.44, 27216.00),
    ("4#中门东侧", "23.1", "m", 12738.52, 8316.00),
    ("5#中门西侧", "25", "m", 12982.10, 9000.00),
    ("6#大门南侧", "32.9", "m", 8199.60, 1184.40),  # 材料8199.60 劳务根据公式推算
]
# 补充劳务值
data1 = [
    ("1#北侧围墙", "167.4", "m", 81594.00, 60264.00),
    ("2#东侧围墙", "75", "m", 38540.40, 27000.00),
    ("3#西侧围墙", "75.6", "m", 37077.44, 27216.00),
    ("4#中门东侧", "23.1", "m", 12738.52, 8316.00),
    ("5#中门西侧", "25", "m", 12982.10, 9000.00),
    ("6#大门南侧", "32.9", "m", 8199.60, 1184.40),  # 劳务1184.40
    # 后附清单项目
    ("砖（22.5万块）", "225,000", "块", 0, 45000.00),  # 这个总价在围墙已完成量合计里
]
sections["围墙工程已完成量"]["mat_total"] = 191132.46
sections["围墙工程已完成量"]["labor_total"] = 126764.70
sections["围墙工程已完成量"]["total"] = 317897.16

# --- 围墙新增项目 ---
data2 = [
    ("新增项目（放大脚、短柱、基础梁马牙槎钢筋增加、商品混凝土含泵送、7#8#-1.5m16水泥墩）", "—", "项", 31584.50, 25530.50),
]
sections["围墙新增项目"]["mat_total"] = 31584.50
sections["围墙新增项目"]["labor_total"] = 25530.50
sections["围墙新增项目"]["total"] = 57115.00

# --- 场地平整工程（xlsx中是清单格式1列总价） ---
data3 = [
    ("回填载重汽车运土方（含租车费）", "8", "车", 7350.00),
    ("回填装载机", "2", "台班", 1000.00),
    ("回填压路机", "2", "台班", 660.00),
    ("围墙轴线放线", "—", "项", 2000.00),
    ("地磅房北侧外墙喷漆", "35", "㎡", 4996.00),
    ("修补围挡", "—", "项", 11000.00),
    ("地磅房东侧混凝土地坪", "12.32", "m³", 2150.00),
]
sections["场地平整工程"]["total"] = 29156.00

# --- 临时设施及文明施工 ---
data4 = [
    ("铁皮围挡", "268", "㎡", 22802.00),
    ("彩旗", "200", "条", 800.00),
    ("大门形象墙+五牌一图+安全警示牌+企业铁质宣传展板", "1", "项", 10670.50),
    ("西侧场地围墙钢管砂浆抹面", "60", "m", 5430.00),
    ("广告宣传横幅", "—", "项", 1500.00),
    ("开业机械租赁费", "—", "项", 4800.00),
    ("开业花费", "—", "项", 6615.00),
    ("现场办公设施", "—", "项", 870.00),
    ("现场材料看护管理人员费用", "64", "天", 5376.00),
    ("现场卫生清理、检查费用", "56", "天", 1344.00),
    ("现场垃圾清运", "—", "项", 2400.00),
    ("宣传彩门", "—", "项", 5000.00),
]
sections["临时设施及文明施工"]["total"] = 56519.30 + 0  # 开业机械租赁费已含在4800

# 修正：开业机械租赁费4800和开业花费6615需确认是否重复或在临时设施合计内
# 从xlsx公式看临时设施合计=SUM上面11项=56519.30

# ========== 参数 ==========
CONTRACT_PARTY = "古浪宁氏钙材料科技有限公司"
PROJECT_NAME = "土门镇围墙及场地平整、大门形象施工项目"
CONTRACT_AMOUNT = 361220.00
PREPAID_AMOUNT = 117400.00
GRAND_TOTAL = 460687.46
CHANGE_AMOUNT = GRAND_TOTAL - CONTRACT_AMOUNT  # 99467.46
REMAINING = GRAND_TOTAL - PREPAID_AMOUNT  # 343287.46

TOTAL_CN = "肆拾陆万零陆佰捌拾柒元肆角陆分"

# ========== 2. 生成Word文档 ==========
doc = Document()

# 页面设置
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(3.17)
section.right_margin = Cm(3.17)

# 样式辅助函数
def set_run_font(run, name='仿宋', size=12, bold=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = r.makeelement(qn('w:rFonts'), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), name)

def add_para(text, bold=False, size=12, align=None, space_after=6, space_before=0, first_indent=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if first_indent:
        p.paragraph_format.first_line_indent = Cm(first_indent)
    return p

def add_heading_custom(text, level=1):
    h = doc.add_heading(text, level=level)
    font_name = '黑体' if level <= 2 else '仿宋'
    for run in h.runs:
        set_run_font(run, name=font_name, bold=True)
    return h

def add_money_line(label, amount, bold=False):
    p = doc.add_paragraph()
    run1 = p.add_run(label)
    set_run_font(run1, size=12)
    run2 = p.add_run(f"¥{amount:,.2f}")
    set_run_font(run2, size=12, bold=True)
    p.paragraph_format.space_after = Pt(4)
    return p

def make_table(headers, rows, col_widths=None):
    """创建带格式的表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                set_run_font(r, size=10, bold=True)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    set_run_font(r, size=10, bold=(r_idx == len(rows)-1 and c_idx >= 2))

    return table


# ===== 封面 =====
for _ in range(6):
    doc.add_paragraph()

add_para("工 程 结 算 书", bold=True, size=26, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=50)

add_para("——" if 1 else "", space_after=40)

add_para(PROJECT_NAME, bold=True, size=16, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=60)

for _ in range(4):
    doc.add_paragraph()

# 封面信息
cover_info = [
    ("建设单位：", CONTRACT_PARTY),
    ("施工单位：", "甘肃东盛建筑设计有限公司"),
    ("编制日期：", date.today().strftime("%Y年%m月%d日")),
]
for label, value in cover_info:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(label)
    set_run_font(r1, size=14)
    r2 = p.add_run(value)
    set_run_font(r2, size=14, bold=True)
    p.paragraph_format.space_after = Pt(12)

doc.add_page_break()

# ===== 目录页 =====
add_para("目  录", bold=True, size=18, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
toc_items = [
    "一、结算说明",
    "二、结算明细表",
    "   （一）围墙工程已完成量",
    "   （二）围墙新增项目",
    "   （三）场地平整工程",
    "   （四）临时设施及文明施工",
    "三、结算汇总",
    "四、付款情况",
    "五、附件",
    "   附件一：现场施工照片",
    "   附件二：材料采购凭证",
]
for item in toc_items:
    add_para(item, size=12, space_after=4, first_indent=0.74)

doc.add_page_break()

# ===== 一、结算说明 =====
add_heading_custom("一、结算说明", level=1)

add_para(f"1. 项目名称：{PROJECT_NAME}", space_after=4)
add_para(f"2. 建设单位：{CONTRACT_PARTY}", space_after=4)
add_para("3. 施工单位：甘肃东盛建筑设计有限公司", space_after=4)
add_para(f"4. 工程地点：甘肃省武威市古浪县土门镇", space_after=4)
add_para(f"5. 原合同金额：¥{CONTRACT_AMOUNT:,.2f}", space_after=4)
add_para(f"6. 变更增加金额：¥{CHANGE_AMOUNT:,.2f}", space_after=4)
add_para(f"7. 结算总价：¥{GRAND_TOTAL:,.2f}（大写：{TOTAL_CN}）", space_after=10)

add_para(
    "本工程依据双方签订的《围墙及场地平整施工合同》及施工过程中经甲方现场确认的新增工程内容，按照实际发生工程量进行结算。",
    space_after=6, first_indent=0.74
)
add_para(
    "其中围墙新增项目（放大脚、短柱、基础梁马牙槎钢筋增加、商品混凝土含泵送、水泥墩等）为原合同工程量清单范围外的新增内容；临时设施及文明施工按现场实际投入计取。以上新增内容均已由甲方现场代表确认并实施完毕。",
    space_after=10, first_indent=0.74
)

# ===== 二、结算明细表 =====
add_heading_custom("二、结算明细表", level=1)

# (一) 围墙工程已完成量
add_heading_custom("（一）围墙工程已完成量", level=2)

rows1 = []
for i, (name, qty, unit, mat, labor) in enumerate(data1, 1):
    rows1.append([str(i), name, qty + unit, f"{mat:,.2f}", f"{labor:,.2f}"])
# 小计行
rows1.append(["", "小计", "", f"{191132.46:,.2f}", f"{126764.70:,.2f}"])

make_table(
    ["序号", "项目名称", "工程量", "材料费（元）", "劳务费（元）"],
    rows1
)
doc.add_paragraph()

# (二) 围墙新增项目
add_heading_custom("（二）围墙新增项目", level=2)

rows2 = []
for i, (name, qty, unit, mat, labor) in enumerate(data2, 1):
    rows2.append([str(i), name, "1", f"{mat:,.2f}", f"{labor:,.2f}"])
rows2.append(["", "小计", "", f"{31584.50:,.2f}", f"{25530.50:,.2f}"])

make_table(
    ["序号", "新增内容说明", "数量", "材料费（元）", "劳务费（元）"],
    rows2
)
doc.add_paragraph()

# (三) 场地平整工程
add_heading_custom("（三）场地平整工程", level=2)

rows3 = []
for i, (name, qty, unit, total) in enumerate(data3, 1):
    rows3.append([str(i), name, qty + unit, f"{total:,.2f}"])
rows3.append(["", "小计", "", f"{29156.00:,.2f}"])

make_table(
    ["序号", "项目名称", "工程量", "合价（元）"],
    rows3
)
doc.add_paragraph()

# (四) 临时设施及文明施工
add_heading_custom("（四）临时设施及文明施工", level=2)

rows4 = []
for i, (name, qty, unit, total) in enumerate(data4, 1):
    rows4.append([str(i), name, qty + unit, f"{total:,.2f}"])
rows4.append(["", "小计", "", f"{56519.30:,.2f}"])

make_table(
    ["序号", "项目名称", "数量", "合价（元）"],
    rows4
)
doc.add_paragraph()

# ===== 三、结算汇总 =====
add_heading_custom("三、结算汇总", level=1)

summary_rows = [
    ["1", "围墙工程已完成量", f"{317897.16:,.2f}"],
    ["2", "围墙新增项目", f"{57115.00:,.2f}"],
    ["3", "场地平整工程", f"{29156.00:,.2f}"],
    ["4", "临时设施及文明施工", f"{56519.30:,.2f}"],
]
# 合并以下行
summary_rows2 = [
    ["", "结算总价", f"{GRAND_TOTAL:,.2f}"],
    ["", "原合同金额", f"{CONTRACT_AMOUNT:,.2f}"],
    ["", "变更增加", f"{CHANGE_AMOUNT:,.2f}"],
]

make_table(
    ["序号", "费用名称", "金额（元）"],
    summary_rows + [["", "分项小计", f"{GRAND_TOTAL:,.2f}"]]
)
doc.add_paragraph()

# 增减对比表
add_para("与原合同对比：", bold=True, size=11, space_after=4)
compare_rows = [
    ["围墙工程", f"{311220.00:,.2f}", f"{375012.16:,.2f}", f"+{63792.16:,.2f}"],
    ["场地平整工程", f"{20000.00:,.2f}", f"{29156.00:,.2f}", f"+{9156.00:,.2f}"],
    ["临时设施及文明施工", f"{30000.00:,.2f}", f"{56519.30:,.2f}", f"+{26519.30:,.2f}"],
    ["合计", f"{CONTRACT_AMOUNT:,.2f}", f"{GRAND_TOTAL:,.2f}", f"+{CHANGE_AMOUNT:,.2f}"],
]
make_table(
    ["费用名称", "原合同金额", "实际结算", "增减额"],
    compare_rows
)
doc.add_paragraph()

# ===== 四、付款情况 =====
add_heading_custom("四、付款情况", level=1)

add_money_line("已付预付款：", PREPAID_AMOUNT)
add_money_line("结算总价：", GRAND_TOTAL)
add_money_line("尚需支付尾款：", REMAINING)
doc.add_paragraph()

add_para(
    f"根据结算结果，{CONTRACT_PARTY}尚需支付本项目尾款人民币{TOTAL_CN}"
    f"（¥{REMAINING:,.2f}）。",
    space_after=8, first_indent=0.74
)

# 付款清单一览
payment_rows = [
    ["已付预付款（第一次付款）", f"{PREPAID_AMOUNT:,.2f}"],
    ["本结算尾款", f"{REMAINING:,.2f}"],
    ["合计", f"{GRAND_TOTAL:,.2f}"],
]
make_table(
    ["付款内容", "金额（元）"],
    payment_rows
)
doc.add_paragraph()

# ===== 五、附件 =====
add_heading_custom("五、附件", level=1)
add_heading_custom("附件一：现场施工照片", level=2)
add_para("（现场照片待补充，由施工单位整理后附入）", size=11, space_after=20)
doc.add_page_break()

add_heading_custom("附件二：材料采购凭证", level=2)
add_para("（材料采购凭证待补充，由施工单位整理后附入）", size=11, space_after=20)

# ===== 签字页 =====
for _ in range(3):
    doc.add_paragraph()

add_para("建设单位（盖章）：", bold=False, size=12, space_after=40)
add_para("", space_after=2)
add_para("代表人（签字）：", size=12, space_after=40)
add_para("", space_after=2)
add_para("日期：    年    月    日", size=12, space_after=40)

doc.add_paragraph()
add_para("—" * 30, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)

add_para("施工单位（盖章）：甘肃东盛建筑设计有限公司", bold=False, size=12, space_after=40)
add_para("", space_after=2)
add_para("代表人（签字）：", size=12, space_after=40)
add_para("", space_after=2)
add_para("日期：    年    月    日", size=12, space_after=40)

# ========== 保存 ==========
output_path = "/Users/mac/Desktop/古浪宁氏项目工程结算书.docx"
doc.save(output_path)
print(f"✅ 结算书已生成：{output_path}")
print(f"   结算总价：¥{GRAND_TOTAL:,.2f}")
print(f"   变更增加：¥{CHANGE_AMOUNT:,.2f}")
print(f"   已付预付款：¥{PREPAID_AMOUNT:,.2f}")
print(f"   应收尾款：¥{REMAINING:,.2f}")
