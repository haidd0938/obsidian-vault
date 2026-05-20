#!/usr/bin/env python3
"""Extract text from a .docx file."""
import sys
from docx import Document

doc = Document(sys.argv[1])
for para in doc.paragraphs:
    print(para.text)
# Also extract tables
for i, table in enumerate(doc.tables):
    print(f"\n--- 表格 {i+1} ---")
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        print(" | ".join(cells))
